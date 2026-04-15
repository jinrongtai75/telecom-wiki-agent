import uuid
import base64
import io
from typing import List
import fitz  # PyMuPDF
import docx

from app.models import DocumentObject, ObjectType, DocumentFormat, ParseResult, ConfirmStatus, BBox


def _new_id() -> str:
    return f"obj-{uuid.uuid4().hex[:8]}"


def _doc_id() -> str:
    return str(uuid.uuid4())


class Parser:
    def parse(self, file_bytes: bytes, filename: str) -> ParseResult:
        fmt = self._detect_format(filename)
        doc_id = _doc_id()
        if fmt == DocumentFormat.PDF:
            objects, raw = self._parse_pdf(file_bytes)
        elif fmt == DocumentFormat.DOCX:
            objects, raw = self._parse_docx(file_bytes)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {filename}")
        return ParseResult(document_id=doc_id, format=fmt, objects=objects, raw_content=raw)


    def _detect_format(self, filename: str) -> DocumentFormat:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "pdf":
            return DocumentFormat.PDF
        if ext in ("docx", "doc"):
            return DocumentFormat.DOCX
        raise ValueError(f"지원하지 않는 파일 형식입니다: {filename}")

    # ── PDF ──────────────────────────────────────────────────────────────────
    def _parse_pdf(self, file_bytes: bytes):
        import re as _re

        NUMBERING_RE = _re.compile(r"^\d+(\.\d+)*\.?\s")

        def _detect_pdf_heading(content: str, max_size: float, is_bold: bool) -> tuple:
            """(is_heading, level) 반환"""
            # 1) 넘버링 패턴
            m = NUMBERING_RE.match(content)
            if m:
                prefix = m.group(0).strip().rstrip(".")
                depth = prefix.count(".") + 1
                return True, min(depth + 1, 6)
            # 2) 폰트 크기 / 볼드
            if max_size >= 16:
                return True, 1
            if max_size >= 14:
                return True, 2
            if is_bold and len(content) <= 80:
                return True, 3
            return False, 0

        # 1패스: 페이지별 raw 객체 수집 (TEXT는 block 단위 낱개, 병합은 2패스에서)
        raw_items = []  # dict: type, content, page, bbox, metadata

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num, page in enumerate(doc, start=1):
            pw, ph = page.rect.width, page.rect.height
            page_items = []

            table_rects: List[fitz.Rect] = []
            try:
                for tab in page.find_tables().tables:
                    table_rects.append(fitz.Rect(tab.bbox))
            except Exception:
                pass

            def _inside_table(bx0, by0, bx1, by1) -> bool:
                cx, cy = (bx0 + bx1) / 2, (by0 + by1) / 2
                return any(r.contains(fitz.Point(cx, cy)) for r in table_rects)

            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] == 0:
                    lines = []
                    max_size = 0.0
                    is_bold = False
                    x0_vals = []
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span["text"].strip()
                            if text:
                                lines.append(text)
                                if span["size"] > max_size:
                                    max_size = span["size"]
                                if span["flags"] & 2**4:
                                    is_bold = True
                                x0_vals.append(span["bbox"][0])
                    content = " ".join(lines).strip()
                    if not content:
                        continue
                    b = block["bbox"]
                    if _inside_table(b[0], b[1], b[2], b[3]):
                        continue
                    block_x0 = min(x0_vals) if x0_vals else b[0]
                    is_heading, level = _detect_pdf_heading(content, max_size, is_bold)
                    page_items.append({
                        "type": "text",
                        "content": content,
                        "page": page_num,
                        "bbox": (b[0], b[1], b[2], b[3], pw, ph),
                        "metadata": {"font_size": max_size, "bold": is_bold, "x0": block_x0},
                        "is_heading": is_heading,
                        "level": level,
                    })

                elif block["type"] == 1:
                    b = block["bbox"]
                    bw = b[2] - b[0]
                    bh = b[3] - b[1]
                    # 점선·구분선 필터: 높이 4px 미만이거나 면적 500px² 미만이거나
                    # 종횡비가 극단적으로 얇은 경우(높이/너비 < 0.03) 장식용 선으로 간주하여 제외
                    if bh < 4 or (bw * bh) < 500 or (bw > 0 and bh / bw < 0.03):
                        continue
                    # block에서 xref를 직접 추출. 실패하면 bbox만 기록하고 content는 비워둠
                    # (나중에 _ensure_image_content가 bbox 크롭으로 복원)
                    content = ""
                    try:
                        # PyMuPDF block type=1의 image dict에서 xref 추출
                        img_info = block.get("image")
                        if isinstance(img_info, dict):
                            xref = img_info.get("xref")
                        else:
                            # block["image"]가 없는 경우 block의 번호로 직접 조회
                            xref = None
                            for img in page.get_images(full=True):
                                # img = (xref, smask, w, h, bpc, cs, alt_cs, name, filter, ...)
                                img_rect = page.get_image_rects(img[0])
                                if img_rect and fitz.Rect(img_rect[0]).intersects(fitz.Rect(b)):
                                    xref = img[0]
                                    break
                        if xref:
                            base_image = doc.extract_image(xref)
                            img_bytes = base_image["image"]
                            b64 = base64.b64encode(img_bytes).decode()
                            ext = base_image.get("ext", "png")
                            content = f"data:image/{ext};base64,{b64}"
                    except Exception:
                        content = ""
                    page_items.append({
                        "type": "image",
                        "content": content,
                        "page": page_num,
                        "bbox": (b[0], b[1], b[2], b[3], pw, ph),
                        "metadata": {},
                    })

            try:
                for tab in page.find_tables().tables:
                    rows = tab.extract()
                    md_rows = []
                    for i, row in enumerate(rows):
                        cells = [str(c or "").replace("\n", " ") for c in row]
                        md_rows.append("| " + " | ".join(cells) + " |")
                        if i == 0:
                            md_rows.append("|" + "|".join(["---"] * len(cells)) + "|")
                    content = "\n".join(md_rows)
                    tb = tab.bbox
                    page_items.append({
                        "type": "table",
                        "content": content,
                        "page": page_num,
                        "bbox": (tb[0], tb[1], tb[2], tb[3], pw, ph),
                        "metadata": {},
                    })
            except Exception:
                pass

            # y0 정렬
            page_items.sort(key=lambda o: o["bbox"][1])
            raw_items.extend(page_items)

        doc.close()

        # 2패스: heading 사이 본문 병합 + 들여쓰기 추정
        # 전체 본문 x0 분포로 기준 x0(좌측 마진) 계산
        body_x0s = [
            it["metadata"]["x0"]
            for it in raw_items
            if it["type"] == "text" and not it.get("is_heading") and it["metadata"].get("x0") is not None
        ]
        base_x0 = sorted(body_x0s)[len(body_x0s) // 4] if len(body_x0s) >= 4 else (min(body_x0s) if body_x0s else 0)
        INDENT_STEP = 15  # px 단위, 이 이상 들여쓰면 1레벨 추가

        def _estimate_ilvl(x0: float) -> int:
            diff = x0 - base_x0
            if diff < INDENT_STEP:
                return 0
            return min(int(diff // INDENT_STEP), 4)

        objects: List[DocumentObject] = []
        order = 0
        body_buf: list = []   # (indent_text, bbox_first)

        def _flush_body():
            nonlocal order
            if not body_buf:
                return
            merged = "\n".join(t for t, _ in body_buf)
            # tuple: (x0, y0, x1, y1, pw, ph, page)
            fb = body_buf[0][1]
            obj = DocumentObject(
                id=_new_id(),
                type=ObjectType.TEXT,
                content=merged,
                order=order,
                page=fb[6],
                bbox=BBox(x0=fb[0], y0=fb[1], x1=fb[2], y1=fb[3], page_width=fb[4], page_height=fb[5]),
                metadata={},
            )
            objects.append(obj)
            body_buf.clear()
            order += 1

        for it in raw_items:
            if it["type"] == "text":
                if it.get("is_heading"):
                    _flush_body()
                    b = it["bbox"]
                    obj = DocumentObject(
                        id=_new_id(),
                        type=ObjectType.TEXT,
                        content=it["content"],
                        order=order,
                        page=it["page"],
                        bbox=BBox(x0=b[0], y0=b[1], x1=b[2], y1=b[3], page_width=b[4], page_height=b[5]),
                        metadata={"font_size": it["metadata"].get("font_size", 0),
                                  "bold": it["metadata"].get("bold", False),
                                  "tag": f"h{it['level']}"},
                        is_heading=True,
                    )
                    objects.append(obj)
                    order += 1
                else:
                    b = it["bbox"]
                    x0 = it["metadata"].get("x0", b[0])
                    ilvl = _estimate_ilvl(x0)
                    indent_text = "  " * ilvl + "- " + it["content"] if ilvl > 0 else it["content"]
                    body_buf.append((indent_text, (b[0], b[1], b[2], b[3], b[4], b[5], it["page"])))

            elif it["type"] == "image":
                _flush_body()
                b = it["bbox"]
                obj = DocumentObject(
                    id=_new_id(),
                    type=ObjectType.IMAGE,
                    content=it["content"],
                    order=order,
                    page=it["page"],
                    bbox=BBox(x0=b[0], y0=b[1], x1=b[2], y1=b[3], page_width=b[4], page_height=b[5]),
                    metadata={},
                )
                objects.append(obj)
                order += 1

            elif it["type"] == "table":
                _flush_body()
                b = it["bbox"]
                obj = DocumentObject(
                    id=_new_id(),
                    type=ObjectType.TABLE,
                    content=it["content"],
                    order=order,
                    page=it["page"],
                    bbox=BBox(x0=b[0], y0=b[1], x1=b[2], y1=b[3], page_width=b[4], page_height=b[5]),
                    metadata={},
                )
                objects.append(obj)
                order += 1

        _flush_body()
        return objects, ""

    # ── DOCX ─────────────────────────────────────────────────────────────────
    def _parse_docx(self, file_bytes: bytes):
        from lxml import etree as _etree
        import re as _re

        objects: List[DocumentObject] = []
        html_parts: List[str] = []
        order = 0

        document = docx.Document(io.BytesIO(file_bytes))
        heading_styles = {"heading 1", "heading 2", "heading 3", "heading 4", "heading 5", "heading 6"}

        # 이미지 rId -> base64 매핑 미리 구성
        img_map: dict = {}
        for rId, rel in document.part.rels.items():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    ext = rel.target_part.content_type.split("/")[-1] or "png"
                    b64 = base64.b64encode(img_bytes).decode()
                    img_map[rId] = f"data:image/{ext};base64,{b64}"
                except Exception:
                    pass

        # Word XML 네임스페이스
        WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        DNS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        ANS = "http://schemas.openxmlformats.org/drawingml/2006/main"
        BLIP_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

        def _para_has_image(para) -> list:
            """단락 내 이미지 rId 목록 반환"""
            found = []
            for drawing in para._element.iter(f"{{{DNS}}}inline", f"{{{DNS}}}anchor"):
                for blip in drawing.iter(f"{{{BLIP_NS}}}blip"):
                    rId = blip.get(f"{{{RNS}}}embed")
                    if rId and rId in img_map:
                        found.append(rId)
            return found

        def _get_list_ilvl(para) -> int:
            """numbering 단락이면 ilvl(0-based) 반환, 아니면 -1"""
            pPr = para._element.pPr
            if pPr is None:
                return -1
            numPr = pPr.find(f"{{{WNS}}}numPr")
            if numPr is None:
                return -1
            numId_el = numPr.find(f"{{{WNS}}}numId")
            if numId_el is None:
                return -1
            numId = numId_el.get(f"{{{WNS}}}val", "0")
            if numId == "0":
                return -1
            ilvl_el = numPr.find(f"{{{WNS}}}ilvl")
            ilvl = int(ilvl_el.get(f"{{{WNS}}}val", "0")) if ilvl_el is not None else 0
            return ilvl

        def _table_to_md(table) -> str:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows.append("|" + "|".join(["---"] * len(cells)) + "|")
            return "\n".join(rows)

        def _table_to_html(table, obj_id: str) -> str:
            tbl = [f"<table data-obj-id='{obj_id}' border='1' cellpadding='4' style='border-collapse:collapse;font-size:13px;cursor:pointer'>"]
            for i, row in enumerate(table.rows):
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                tag = "th" if i == 0 else "td"
                tbl.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
            tbl.append("</table>")
            return "\n".join(tbl)

        def _escape(s: str) -> str:
            return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # 넘버링 패턴: "1.", "5.1", "5.1.1", "3.2.1.4 " 등
        NUMBERING_RE = _re.compile(r"^\d+(\.\d+)*\.?\s")

        def _detect_heading(para) -> tuple[bool, int]:
            """(is_heading, level) 반환. level: 1~6"""
            style_name = para.style.name.lower() if para.style else ""
            content = para.text.strip()

            # 1) Word 스타일 기반
            if style_name in heading_styles:
                m = _re.search(r"(\d)", style_name)
                level = min(int(m.group(1)) if m else 2, 6)
                return True, level

            # 2) 넘버링 패턴 기반
            m = NUMBERING_RE.match(content)
            if m:
                # 점의 개수로 depth 계산 (예: "5." → 1, "5.1" → 2, "5.1.1" → 3)
                prefix = m.group(0).strip().rstrip(".")
                depth = prefix.count(".") + 1
                level = min(depth + 1, 6)  # depth 1 → h2, depth 2 → h3, ...
                return True, level

            # 3) 볼드 단독 짧은 줄 (50자 이하)
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            if is_bold and len(content) <= 50:
                return True, 3

            return False, 0

        # 본문 단락 버퍼 flush → 하나의 TEXT 객체로 병합
        body_buf: list[str] = []   # 텍스트 라인
        body_html_buf: list[str] = []  # HTML 라인

        def _flush_body():
            nonlocal order
            if not body_buf:
                return
            merged = "\n".join(body_buf)
            obj_id = _new_id()
            obj = DocumentObject(
                id=obj_id,
                type=ObjectType.TEXT,
                content=merged,
                order=order,
                metadata={},
                is_heading=False,
            )
            objects.append(obj)
            inner_html = "\n".join(body_html_buf)
            html_parts.append(f"<div data-obj-id='{obj_id}' style='cursor:pointer;margin-bottom:8px'>{inner_html}</div>")
            body_buf.clear()
            body_html_buf.clear()
            order += 1

        # body 자식 요소를 XML 순서대로 순회
        body = document.element.body
        para_map = {p._element: p for p in document.paragraphs}
        table_map = {t._element: t for t in document.tables}

        for child in body:
            tag = _etree.QName(child.tag).localname if child.tag else ""

            if tag == "p":
                para = para_map.get(child)
                if para is None:
                    continue

                # 이미지가 있는 단락
                img_rids = _para_has_image(para)
                if img_rids:
                    _flush_body()
                    for rId in img_rids:
                        src = img_map[rId]
                        obj_id = _new_id()
                        obj = DocumentObject(
                            id=obj_id,
                            type=ObjectType.IMAGE,
                            content=src,
                            order=order,
                            metadata={},
                        )
                        objects.append(obj)
                        html_parts.append(f"<p data-obj-id='{obj_id}' style='cursor:pointer'><img src='{src}' style='max-width:100%;pointer-events:none' /></p>")
                        order += 1
                    continue

                content = para.text.strip()
                if not content:
                    continue

                is_heading, level = _detect_heading(para)

                if is_heading:
                    # heading 만나면 본문 버퍼 먼저 flush
                    _flush_body()
                    obj_id = _new_id()
                    obj = DocumentObject(
                        id=obj_id,
                        type=ObjectType.TEXT,
                        content=content,
                        order=order,
                        metadata={"tag": f"h{level}"},
                        is_heading=True,
                    )
                    objects.append(obj)
                    escaped = _escape(content)
                    html_parts.append(f"<h{level} data-obj-id='{obj_id}' style='cursor:pointer'>{escaped}</h{level}>")
                    order += 1
                else:
                    # 본문 단락 → 버퍼에 누적
                    ilvl = _get_list_ilvl(para)
                    escaped = _escape(content)
                    if ilvl >= 0:
                        # 다단계 목록: content에 들여쓰기 prefix
                        indent_text = "  " * ilvl + "- " + content
                        indent_html = f"<p style='margin:2px 0;padding-left:{ilvl * 20}px'>{'- ' + escaped}</p>"
                    else:
                        indent_text = content
                        indent_html = f"<p style='margin:2px 0'>{escaped}</p>"
                    body_buf.append(indent_text)
                    body_html_buf.append(indent_html)

            elif tag == "tbl":
                _flush_body()
                table = table_map.get(child)
                if table is None:
                    continue
                md_content = _table_to_md(table)
                obj_id = _new_id()
                obj = DocumentObject(
                    id=obj_id,
                    type=ObjectType.TABLE,
                    content=md_content,
                    order=order,
                    metadata={},
                )
                objects.append(obj)
                html_parts.append(_table_to_html(table, obj_id))
                order += 1

        # 마지막 본문 버퍼 flush
        _flush_body()

        raw_html = "\n".join(html_parts)
        return objects, raw_html

